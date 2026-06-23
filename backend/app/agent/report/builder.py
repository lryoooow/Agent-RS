"""对话分析报告生成器（Word/.docx）。

对话工具 generate_report 与 HTTP 端点 /api/reports 共用此模块：
读取本对话**真实持久化**的分析结果（list_recent_analysis_results）+ 影像元信息，
用 python-docx 写成可下载的 .docx，落到该影像的 results 目录、复用既有下载路由。

硬约束：报告内容只来自真实结果，无结果即抛 ReportError 拒绝生成、绝不编造（延续
助手"不杜撰数据"的正确行为，并满足项目约束"影像分析结论必须标注数据来源与时间"）。
"""
from __future__ import annotations

import asyncio
import logging
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from app.agent.imagery_access import read_imagery_metadata
from app.core.paths import imagery_root
from app.db.pool import fetch_optional_pool
from app.db.repositories.conversation import get_conversation
from app.db.repositories.message import list_recent_analysis_results

logger = logging.getLogger(__name__)

REPORT_ANALYSIS_LIMIT = 20


class ReportError(Exception):
    """报告无法生成（无持久化结果、非属主、存储不可用等）。文案对用户安全，不含内部细节。"""

    def __init__(self, message: str, *, code: str) -> None:
        super().__init__(message)
        self.code = code
        self.message = message


@dataclass(frozen=True)
class ReportArtifact:
    imagery_id: str
    filename: str
    download_url: str


async def build_conversation_report(
    *,
    conversation_id: str | None,
    user_id: str | None,
    imagery_id: str | None = None,
) -> ReportArtifact:
    """读本对话真实分析结果，生成 .docx 报告，返回下载信息。

    流程：归属校验 → 取持久化分析结果 → 选定影像 → 读影像元信息 →
    （无结果即拒绝）→ to_thread 写 docx → 落 results 目录 → 返回 download_url。
    """
    if not user_id:
        raise ReportError("无法确认用户身份，无法生成报告。", code="no_user")
    if not conversation_id:
        raise ReportError("缺少对话上下文，无法定位要汇总的分析结果。", code="no_conversation")

    pool = await fetch_optional_pool()
    if pool is None:
        raise ReportError("存储未启用，无法读取分析结果生成报告。", code="storage_inactive")

    async with pool.acquire() as conn:
        if await get_conversation(conn, conversation_id, user_id) is None:
            # 非属主或对话不存在：拒绝，不泄漏他人结果。
            raise ReportError("未找到该对话或无权访问，无法生成报告。", code="conversation_forbidden")
        analyses = await list_recent_analysis_results(
            conn,
            conversation_id=conversation_id,
            user_id=user_id,
            limit=REPORT_ANALYSIS_LIMIT,
        )

    selected_id, selected_analyses = _select_imagery_analyses(analyses, imagery_id)
    if not selected_analyses:
        # 没有任何真实分析结果可写——绝不编造空报告。
        raise ReportError(
            "本对话还没有可用于报告的真实分析结果，请先对影像执行分析（如地物分类、目标检测）。",
            code="no_analysis",
        )

    imagery_meta = read_imagery_metadata(selected_id) if selected_id else None
    generated_at = datetime.now(timezone.utc)
    results_dir = imagery_root() / selected_id / "results"

    try:
        results_dir.mkdir(parents=True, exist_ok=True)
        filename = f"report_{generated_at.strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.docx"
        output_path = results_dir / filename
        # docx 写入是同步 CPU/IO，offload 到线程，避免阻塞事件循环（与上传/解析一致）。
        await asyncio.to_thread(
            _render_report_docx,
            output_path,
            imagery_id=selected_id,
            imagery_meta=imagery_meta,
            analyses=selected_analyses,
            generated_at=generated_at,
        )
    except ReportError:
        raise
    except Exception as exc:
        logger.exception("Report generation failed: %s", exc)
        raise ReportError("报告生成失败，请稍后重试。", code="render_failed") from exc

    return ReportArtifact(
        imagery_id=selected_id,
        filename=filename,
        download_url=f"/api/imagery/{selected_id}/results/{filename}",
    )


def _select_imagery_analyses(
    analyses: list[dict[str, Any]],
    imagery_id: str | None,
) -> tuple[str | None, list[dict[str, Any]]]:
    """从对话分析结果里选定一张影像及其全部结果。

    指定 imagery_id → 只取该影像；否则取**最近一次被分析的影像**（analyses 已按时间正序，
    取最后一条出现的 imagery_id）。返回 (影像ID, 该影像的分析结果列表)。
    """
    def _imagery_of(entry: dict[str, Any]) -> str | None:
        for key in ("geospatial_result", "tool_result"):
            payload = entry.get(key)
            if isinstance(payload, dict) and payload.get("imagery_id"):
                return str(payload["imagery_id"])
        return None

    if imagery_id:
        target = imagery_id
    else:
        target = None
        for entry in analyses:  # 正序遍历，最后命中的即最近一张
            found = _imagery_of(entry)
            if found:
                target = found
    if not target:
        return None, []
    selected = [entry for entry in analyses if _imagery_of(entry) == target]
    return target, selected


def _render_report_docx(
    output_path: Path,
    *,
    imagery_id: str,
    imagery_meta: dict[str, Any] | None,
    analyses: list[dict[str, Any]],
    generated_at: datetime,
) -> None:
    """用 python-docx 写报告正文（纯同步，被 to_thread 调用）。

    只渲染真实存在的字段：影像元信息缺失则跳过该行，分析结果缺字段则不写。
    末尾固定标注数据来源与生成时间（满足项目约束）。
    """
    from docx import Document

    document = Document()
    document.add_heading("遥感影像分析报告", level=0)

    ts_text = generated_at.strftime("%Y-%m-%d %H:%M UTC")
    document.add_paragraph(f"影像 ID：{imagery_id}")
    document.add_paragraph(f"生成时间：{ts_text}")

    # 影像基本信息表（仅写出 metadata 里真实存在的字段）
    document.add_heading("一、影像基本信息", level=1)
    meta = imagery_meta or {}
    info_rows = [
        ("原始文件名", meta.get("filename")),
        ("坐标系 (CRS)", meta.get("crs")),
        ("尺寸", f"{meta.get('width')}×{meta.get('height')} px" if meta.get("width") and meta.get("height") else None),
        ("波段数", meta.get("band_count")),
        ("数据类型", meta.get("dtype")),
    ]
    present_rows = [(label, value) for label, value in info_rows if value not in (None, "")]
    if present_rows:
        table = document.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for label, value in present_rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = str(value)
    else:
        document.add_paragraph("（影像元信息缺失）")

    # 各分析结果
    document.add_heading("二、分析结果", level=1)
    section_no = 0
    for entry in analyses:
        section_no += _render_analysis_entry(document, entry, section_no)
    if section_no == 0:
        document.add_paragraph("（无可呈现的结构化分析结果）")

    # 数据来源与边界声明
    document.add_heading("三、数据来源与说明", level=1)
    document.add_paragraph(
        "本报告所有数值均来自本对话中对上述影像实际执行的工具计算结果（地物分类、"
        "目标检测、光谱指数等），未做任何人工编造或估算。各模型结论存在适用边界："
        "地物分类基于 LandCover.ai 四类体系，目标检测基于 DOTA 15 类，"
        "指数解读阈值随传感器、地区、季节、大气校正变化，引用时请结合实际场景判断。"
    )
    document.add_paragraph(f"报告生成时间：{ts_text}")

    document.save(str(output_path))


def _render_analysis_entry(document, entry: dict[str, Any], current_no: int) -> int:
    """渲染一条分析结果到文档；成功渲染返回 1，无法识别返回 0。"""
    geo = entry.get("geospatial_result")
    tool = entry.get("tool_result")
    if isinstance(geo, dict):
        geo_type = geo.get("type")
        if geo_type == "segmentation":
            _render_segmentation(document, geo, current_no + 1)
            return 1
        if geo_type == "detection":
            _render_detection(document, geo, current_no + 1)
            return 1
        if geo_type in ("ndvi", "spectral_index"):
            _render_index(document, geo, current_no + 1)
            return 1
    if isinstance(tool, dict) and tool.get("type") == "raster_inspect":
        _render_raster_inspect(document, tool, current_no + 1)
        return 1
    return 0


def _render_segmentation(document, geo: dict[str, Any], no: int) -> None:
    document.add_heading(f"{no}. 地物分类（LandCover.ai 四类）", level=2)
    classes = [c for c in (geo.get("classes") or []) if isinstance(c, dict)]
    if not classes:
        document.add_paragraph("未识别到地物类别。")
        return
    table = document.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text, header[1].text, header[2].text = "类别", "像素数", "占比"
    for item in classes:
        cells = table.add_row().cells
        cells[0].text = str(item.get("label") or item.get("name") or "未知")
        cells[1].text = str(item.get("pixel_count", "—"))
        pct = item.get("percentage")
        cells[2].text = f"{pct:.2f}%" if isinstance(pct, (int, float)) else "—"


def _render_detection(document, geo: dict[str, Any], no: int) -> None:
    document.add_heading(f"{no}. 目标检测（DOTA 15 类）", level=2)
    total = geo.get("detection_count")
    document.add_paragraph(f"检测目标总数：{total if total is not None else '未知'}")
    classes = [c for c in (geo.get("classes") or []) if isinstance(c, dict)]
    if not classes:
        return
    table = document.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text, header[1].text = "类别", "数量"
    for item in classes:
        cells = table.add_row().cells
        cells[0].text = str(item.get("label") or item.get("name") or "未知")
        cells[1].text = str(item.get("count", "—"))


def _render_index(document, geo: dict[str, Any], no: int) -> None:
    stats = geo.get("stats") or {}
    index_type = geo.get("index_type") or stats.get("index_type") or ("NDVI" if geo.get("type") == "ndvi" else "光谱指数")
    document.add_heading(f"{no}. {index_type} 指数统计", level=2)
    table = document.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for key, label in (("min", "最小值"), ("max", "最大值"), ("mean", "均值"), ("std", "标准差")):
        value = stats.get(key)
        if isinstance(value, (int, float)) and not isinstance(value, bool):
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = f"{value:.4g}"


def _render_raster_inspect(document, tool: dict[str, Any], no: int) -> None:
    document.add_heading(f"{no}. 影像质检", level=2)
    bits = []
    if tool.get("width") and tool.get("height"):
        bits.append(f"尺寸 {tool['width']}×{tool['height']} px")
    if tool.get("band_count"):
        bits.append(f"{tool['band_count']} 波段")
    if tool.get("crs"):
        bits.append(f"CRS {tool['crs']}")
    document.add_paragraph("；".join(bits) if bits else "基本信息缺失。")
