from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.core.settings import Settings
from app.db.repositories._pg import conversation as conv_repo
from app.db.repositories._pg import identity as identity_repo
from app.db.repositories._pg import message as msg_repo

from app.agent.report import builder as report_builder
from app.agent.report.builder import ReportError, build_conversation_report

REPO_ROOT = Path(__file__).resolve().parents[3]

# 仅对异步用例标 asyncio（loop_scope=session 以复用 session 级 pg_pool 的 event loop）；
# 末尾的同步 utf8 自检不标，避免被误判为 async（模块级 pytestmark 会误伤它）。
_pg_async = pytest.mark.asyncio(loop_scope="session")


@pytest.fixture
async def pool(pg_pool):
    """复用全局 PG 池（tests/conftest.py 的 pg_pool）。

    本测试原跑在临时 SQLite 文件上；SQLite 退役后改连真实 PostgreSQL。
    沿用 `pool` 这一夹具名 + async（与已验证的 pg_conn 同模式），使下方测试
    签名（pool, settings, ...）零改动。
    """
    return pg_pool


@pytest.fixture
def settings():
    return Settings(_env_file=None)


async def _seed_conversation_with_segmentation(pool, settings, *, imagery_id: str) -> tuple[str, str]:
    """建一个对话并持久化一条真实地物分类结果，返回 (conversation_id, owner_id)。"""
    async with pool.acquire() as conn:
        await identity_repo.ensure_default_identity(conn, settings)
        owner = settings.default_user_id
        cid = await conv_repo.create_conversation(conn, user_id=owner, settings=settings, title="报告对话")
        await msg_repo.append_message(
            conn, conversation_id=cid, role="assistant", content="地物分类完成",
            metadata={
                "finish_reason": "stop",
                "geospatial_result": {
                    "type": "segmentation",
                    "imagery_id": imagery_id,
                    "total_pixels": 1000,
                    "classes": [
                        {"label": "背景", "pixel_count": 913, "percentage": 91.307},
                        {"label": "建筑", "pixel_count": 34, "percentage": 3.441},
                    ],
                },
            },
        )
    return cid, owner


@_pg_async
async def test_build_report_from_real_results_keeps_readable_chinese(pool, settings, monkeypatch, tmp_path):
    # 常规 + 历史重复点：基于真实持久化分类结果生成 .docx，读回校验中文无乱码、含真实占比与来源声明。
    # 取代旧的"读缺失手工 .docx"烟测，根除既有无关失败。
    imagery_id = "d722c20e1234"
    imagery_dir = tmp_path / "imagery" / imagery_id
    imagery_dir.mkdir(parents=True)
    (imagery_dir / "metadata.json").write_text(
        '{"filename":"GF2_test.tif","crs":"EPSG:4326","width":256,"height":256,'
        '"band_count":4,"dtype":"uint16","owner_user_id":"' + settings.default_user_id + '"}',
        encoding="utf-8",
    )
    monkeypatch.setattr(report_builder, "imagery_root", lambda: tmp_path / "imagery")
    # read_imagery_metadata 走的是 imagery_access 模块内的 imagery_root，需一并指向 tmp。
    monkeypatch.setattr("app.agent.imagery_access.imagery_root", lambda: tmp_path / "imagery")
    monkeypatch.setattr("app.agent.report.builder.fetch_optional_pool", lambda: _wrap(pool))

    cid, _owner = await _seed_conversation_with_segmentation(pool, settings, imagery_id=imagery_id)

    artifact = await build_conversation_report(conversation_id=cid, user_id=settings.default_user_id)

    assert artifact.imagery_id == imagery_id
    assert artifact.filename.endswith(".docx")
    assert artifact.download_url == f"/api/imagery/{imagery_id}/results/{artifact.filename}"

    out = tmp_path / "imagery" / imagery_id / "results" / artifact.filename
    assert out.exists()
    document = Document(str(out))
    text = "\n".join(p.text for p in document.paragraphs)
    table_text = "\n".join(c.text for t in document.tables for r in t.rows for c in r.cells)

    assert "遥感影像分析报告" in text
    assert imagery_id in text
    assert "数据来源" in text and "未做任何人工编造" in text  # 项目约束：标注来源、不编造
    assert "91.31%" in table_text  # 真实占比
    assert "GF2_test.tif" in table_text
    assert chr(0xFFFD) not in (text + table_text)  # 无乱码（U+FFFD 替换符，用 chr 避免本文件触发 utf8 自检）
    assert "？？" not in text


@_pg_async
async def test_build_report_without_results_refuses(pool, settings, monkeypatch, tmp_path):
    # 异常分支：对话无任何真实分析结果 → 抛 ReportError、不产文件、绝不编造空报告。
    monkeypatch.setattr(report_builder, "imagery_root", lambda: tmp_path / "imagery")
    monkeypatch.setattr("app.agent.report.builder.fetch_optional_pool", lambda: _wrap(pool))
    async with pool.acquire() as conn:
        await identity_repo.ensure_default_identity(conn, settings)
        cid = await conv_repo.create_conversation(
            conn, user_id=settings.default_user_id, settings=settings, title="空对话"
        )
        await msg_repo.append_message(
            conn, conversation_id=cid, role="assistant", content="纯文字",
            metadata={"finish_reason": "stop"},
        )

    with pytest.raises(ReportError) as exc_info:
        await build_conversation_report(conversation_id=cid, user_id=settings.default_user_id)
    assert exc_info.value.code == "no_analysis"


@_pg_async
async def test_build_report_rejects_non_owner(pool, settings, monkeypatch, tmp_path):
    # 隔离：非属主请求 → 拒绝，不泄漏他人对话结果。
    monkeypatch.setattr(report_builder, "imagery_root", lambda: tmp_path / "imagery")
    monkeypatch.setattr("app.agent.report.builder.fetch_optional_pool", lambda: _wrap(pool))
    cid, _owner = await _seed_conversation_with_segmentation(pool, settings, imagery_id="d722c20e1234")

    intruder = "00000000-0000-4000-8000-0000000000ff"
    with pytest.raises(ReportError) as exc_info:
        await build_conversation_report(conversation_id=cid, user_id=intruder)
    assert exc_info.value.code == "conversation_forbidden"


@_pg_async
async def test_build_report_requires_conversation(pool, settings, monkeypatch, tmp_path):
    # 非法输入：缺 conversation_id / user_id → 拒绝。
    monkeypatch.setattr("app.agent.report.builder.fetch_optional_pool", lambda: _wrap(pool))
    with pytest.raises(ReportError) as exc_info:
        await build_conversation_report(conversation_id=None, user_id=settings.default_user_id)
    assert exc_info.value.code == "no_conversation"

    with pytest.raises(ReportError) as exc_info:
        await build_conversation_report(conversation_id="x", user_id=None)
    assert exc_info.value.code == "no_user"


def _wrap(pool):
    """把已就绪的 pool 包成 awaitable，匹配 fetch_optional_pool 的 async 调用约定。"""
    async def _inner():
        return pool
    return _inner()


def test_python_sources_are_utf8_without_replacement_characters() -> None:
    replacement = chr(0xFFFD)  # 用 chr 构造，避免本检测文件自身含 U+FFFD 字面量而误报
    scanned = 0
    for base in (REPO_ROOT / "backend" / "app", REPO_ROOT / "backend" / "tests"):
        for path in base.rglob("*.py"):
            if "__pycache__" in path.parts:
                continue
            text = path.read_text(encoding="utf-8")
            scanned += 1
            assert replacement not in text, path.relative_to(REPO_ROOT).as_posix()
    assert scanned > 0
